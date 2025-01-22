from datetime import datetime
from colors import colorize, VERDE, ROJO, AMARILLO, AZUL, CYAN, MAGENTA, NORMAL, BOLD, UNDERLINE, ITALIC
from openpyxl import load_workbook
from docx import Document


def generar_docx_desde_plantilla(ruta_plantilla, ruta_salida, registros):
    """
    Genera un archivo .docx basado en una plantilla, reemplazando los placeholders.

    Args:
        ruta_plantilla (str): Ruta al archivo de plantilla .docx.
        ruta_salida (str): Ruta donde se guardará el archivo generado.
        reemplazos (dict): Diccionario con los placeholders y sus valores.
    """
    contrato = 0
    for registro in registros:
        # Crear una copia de la plantilla
        contrato +=1 
        doc = Document(ruta_plantilla)
    
        # Reemplazar placeholders en cada párrafo
        for parrafo in doc.paragraphs:
            for placeholder, valor in registro.items():
                if placeholder in parrafo.text:
                    parrafo.text = parrafo.text.replace(placeholder, str(valor))
    
    # Reemplazar placeholders en tablas (opcional)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for placeholder, valor in registro.items():
                    if placeholder in celda.text:
                        celda.text = celda.text.replace(placeholder, str(valor))
    
    # Guardar el archivo generado
    nombre_archivo = f"{ruta_salida}/contrato_{contrato}.docx"
    doc.save(nombre_archivo)
    print(f"Archivo generado en: {ruta_salida}")


def obtener_datos_excel(ruta_archivo):
    """
    Obtiene los datos de una hoja de Excel y los devuelve como un diccionario.
    
    Args:
        ruta_archivo (str): Ruta al archivo de Excel.
    
    Returns:
        dict: Diccionario con los datos de la hoja de Excel.
    """
    wb = load_workbook(ruta_archivo)
    ws = wb['Datos']
    datos = []
    
    headers = [cell.value for cell in next(ws.iter_rows(max_row=1))]

    rif_idx = headers.index("RIF")

    name_idx = headers.index("Razón Social o Nombre")

    phone_idx = headers.index("TELEFONO")

    email_idx = headers.index("CORREO ELECTRÓNICO")
    
    for row in ws.iter_rows(min_row=2, values_only=True):
        registro = {
                "RIF": row[rif_idx],
                "Nombre": row[name_idx],
                "Teléfono": row[phone_idx],
                "Correo": row[email_idx]
            }
        datos.append(registro)
        
    return datos

    return None




# Ejemplo de uso
ruta_plantilla = "Data/Plantillas/Pantilla1.docx"  # Ruta del archivo de plantilla
ruta_salida = "Data/Contratos/"  # Ruta del archivo generado
# registros = obtener_datos_excel("Data/Plantillas/hoja.xlsx")  # Lista de diccionarios con los registros
registros = [
    {
        "{nombre}": "Alexis",
        "{fecha}": "21 de enero de 2025",
        "{mensaje}": "¡Este es un ejemplo generado dinámicamente!"
    }
]

generar_docx_desde_plantilla(ruta_plantilla, ruta_salida, registros)